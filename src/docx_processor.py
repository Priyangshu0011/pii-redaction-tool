import os
import docx
from typing import Dict, Any, List, Tuple
from src.pii_detector import PIIDetector
from src.anonymizer import Anonymizer

class DocxProcessor:
    """
    XML Run-aware Word Document (.docx) PII Redactor.
    Traverses paragraphs, tables, headers, footers, and hyperlinks while
    preserving font styles, bold/italic, colors, and XML structure.
    """

    def __init__(self, detector: PIIDetector = None, anonymizer: Anonymizer = None):
        self.detector = detector or PIIDetector()
        self.anonymizer = anonymizer or Anonymizer()

    def process_document(self, input_path: str, output_path: str, entity_types: List[str] = None) -> Dict[str, Any]:
        """
        Read input docx file, detect and redact PII, write to output_path,
        and return summary statistics.
        """
        if not os.path.exists(input_path):
            raise FileNotFoundError(f"Input file not found: {input_path}")

        doc = docx.Document(input_path)
        stats = {
            "total_entities_detected": 0,
            "entities_by_type": {},
            "paragraphs_processed": 0,
            "tables_processed": 0,
            "headers_footers_processed": 0,
            "detected_list": []
        }

        self.anonymizer.reset()

        # 1. Process Main Paragraphs
        for paragraph in doc.paragraphs:
            count = self._process_paragraph(paragraph, stats, entity_types)
            stats["paragraphs_processed"] += 1

        # 2. Process Tables
        for table in doc.tables:
            self._process_table(table, stats, entity_types)
            stats["tables_processed"] += 1

        # 3. Process Headers and Footers (Deduplicated)
        visited_parts = set()
        for section in doc.sections:
            for part in (section.header, section.first_page_header, section.even_page_header,
                         section.footer, section.first_page_footer, section.even_page_footer):
                if part and part._element not in visited_parts:
                    visited_parts.add(part._element)
                    for hp in part.paragraphs:
                        self._process_paragraph(hp, stats, entity_types)
                    for ht in part.tables:
                        self._process_table(ht, stats, entity_types)

            stats["headers_footers_processed"] += 1

        # Save output document
        doc.save(output_path)
        return stats

    def _process_table(self, table, stats: Dict[str, Any], entity_types: List[str]):
        """Recursively process table cells, deduplicating merged cells."""
        visited_cells = set()
        for row in table.rows:
            for cell in row.cells:
                if cell._tc not in visited_cells:
                    visited_cells.add(cell._tc)
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, stats, entity_types)

    def _process_paragraph(self, paragraph, stats: Dict[str, Any], entity_types: List[str]) -> int:
        """
        Process a single paragraph: detect entities across stitched run text,
        replace text in runs, and update stats.
        """
        if not paragraph.runs or not paragraph.text.strip():
            return 0

        full_text = paragraph.text
        entities = self.detector.detect(full_text, entity_types=entity_types)

        if not entities:
            return 0

        # Sort entities in reverse order (from end to start) so index replacements don't shift earlier offsets
        sorted_entities = sorted(entities, key=lambda x: x["start"], reverse=True)

        for ent in sorted_entities:
            orig_text = ent["text"]
            pii_type = ent["type"]
            replacement = self.anonymizer.anonymize(orig_text, pii_type)

            # Record stats
            stats["total_entities_detected"] += 1
            stats["entities_by_type"][pii_type] = stats["entities_by_type"].get(pii_type, 0) + 1
            stats["detected_list"].append({
                "original": orig_text,
                "replacement": replacement,
                "type": pii_type,
                "source": ent.get("source", "UNKNOWN")
            })

            start_char = ent["start"]
            end_char = ent["end"]

            self._replace_text_in_runs(paragraph, start_char, end_char, replacement)

        return len(entities)

    def _replace_text_in_runs(self, paragraph, start_idx: int, end_idx: int, replacement: str):
        """
        Maps paragraph character offsets [start_idx, end_idx) to individual runs
        and replaces the target substring with replacement text.
        """
        run_boundaries = []
        curr_len = 0
        for run in paragraph.runs:
            r_start = curr_len
            r_end = curr_len + len(run.text)
            run_boundaries.append((run, r_start, r_end))
            curr_len = r_end

        first_affected_run = None
        for run, r_start, r_end in run_boundaries:
            if r_start < end_idx and r_end > start_idx:
                # This run overlaps with [start_idx, end_idx)
                overlap_start = max(start_idx, r_start) - r_start
                overlap_end = min(end_idx, r_end) - r_start

                if first_affected_run is None:
                    first_affected_run = run
                    # Replace overlapped text in first affected run with replacement
                    run.text = run.text[:overlap_start] + replacement + run.text[overlap_end:]
                else:
                    # Clear overlapped portion in subsequent affected runs
                    run.text = run.text[:overlap_start] + run.text[overlap_end:]
